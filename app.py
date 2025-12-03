import io
import re
import html
import zipfile

import requests
import pandas as pd
import streamlit as st

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF  # NEW: for PDF support


# ================== Core Logic (Shared) ================== #

ALEPH_BET = [
    "א", "ב", "ג", "ד", "ה", "ו", "ז", "ח", "ט", "י",
    "כ", "ל", "מ", "נ", "ס", "ע", "פ", "צ", "ק", "ר",
    "ש", "ת",
]

letter_to_index = {letter: i for i, letter in enumerate(ALEPH_BET)}
TAG_RE = re.compile(r"<[^>]+>")


def clean_hebrew_verse(raw: str) -> str:
    """Remove HTML tags, entities, and parsha markers from a verse."""
    text = html.unescape(raw)
    text = TAG_RE.sub("", text)
    text = text.replace("{פ}", "").replace("{ס}", "")
    return text.strip()


@st.cache_data
def load_tehillim_119():
    """
    Fetch Tehillim 119 (Hebrew) from Sefaria, clean it,
    and return a list of 176 verses.
    Cached so it only runs once.
    """
    url = "https://www.sefaria.org/api/texts/Psalms.119?lang=he&context=0"
    resp = requests.get(url)
    resp.raise_for_status()
    data = resp.json()
    verses = data["he"]
    cleaned = [clean_hebrew_verse(v) for v in verses]

    if len(cleaned) != 176:
        raise ValueError(f"Expected 176 verses, got {len(cleaned)}")

    return cleaned


def build_stanzas(verses_119):
    """Split 176 verses into 22 stanzas of 8 pesukim."""
    return [verses_119[i:i + 8] for i in range(0, 176, 8)]


def get_stanzas_for_name(name, stanzas):
    """Return list of (letter, stanza) pairs for a Hebrew name."""
    sections = []

    final_to_regular = {
        "ך": "כ", "ם": "מ", "ן": "נ", "ף": "פ", "ץ": "צ"
    }

    name = name.strip()

    for ch in name:
        if ch == " ":
            continue

        letter = final_to_regular.get(ch, ch)
        idx = letter_to_index.get(letter)

        if idx is not None:
            sections.append((letter, stanzas[idx]))
    return sections


def build_docx_bytes_for_name(name, stanzas):
    """
    Build a DOCX in memory for a given name and return bytes.
    """
    sections = get_stanzas_for_name(name, stanzas)
    if not sections:
        raise ValueError(f"No valid Hebrew letters found in name '{name}'.")

    doc = Document()

    title_para = doc.add_paragraph(f"תהילים פרק קיט עבור השם: {name}")
    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for letter, stanza in sections:
        p_letter = doc.add_paragraph(letter)
        p_letter.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for pasuk in stanza:
            p = doc.add_paragraph(pasuk)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph("").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def build_pdf_bytes_for_name(name, stanzas):
    """
    Build a PDF in memory for a given name and return raw bytes.
    Uses fpdf2 with a Unicode font that supports Hebrew.
    Assumes font/DejaVuSansCondensed.ttf is present.
    """
    sections = get_stanzas_for_name(name, stanzas)
    if not sections:
        raise ValueError(f"No valid Hebrew letters found in name '{name}'.")

    pdf = FPDF()
    # Enable text shaping for RTL languages (requires fpdf2 >= 2.7)
    try:
        pdf.set_text_shaping(True)
    except AttributeError:
        # Older versions may not have set_text_shaping; PDF will still render, just less perfect shaping
        pass

    pdf.add_page()

    # Load Unicode font that supports Hebrew
    # Make sure font/DejaVuSansCondensed.ttf exists in the repo
    pdf.add_font("DejaVu", "", "font/DejaVuSansCondensed.ttf", uni=True)
    pdf.set_font("DejaVu", size=14)
    pdf.set_auto_page_break(auto=True, margin=15)

    # Title
    title = f"תהילים פרק קיט עבור השם: {name}"
    pdf.multi_cell(0, 8, txt=title, align="R")
    pdf.ln(4)

    for letter, stanza in sections:
        # Letter heading
        pdf.multi_cell(0, 8, txt=letter, align="R")
        pdf.ln(1)

        for pasuk in stanza:
            pdf.multi_cell(0, 8, txt=pasuk, align="R")
        pdf.ln(3)

    # Return raw bytes
    pdf_bytes = pdf.output(dest="S").encode("latin1")
    return pdf_bytes


# ================== Streamlit UI ================== #

st.set_page_config(page_title="Perek 119 Builder", page_icon="📖", layout="centered")

st.title("📖 Perek119Builder – Web Version")
st.write(
    "Create Tehillim 119 (פרק קיט) documents based on Hebrew names.\n\n"
    "You can either upload an Excel file with many names, or enter a single name directly."
)

# Load Tehillim 119 once
with st.spinner("Loading Tehillim 119 text..."):
    verses_119 = load_tehillim_119()
    stanzas_119 = build_stanzas(verses_119)

st.divider()

tab_single, tab_excel = st.tabs(["🔹 Single Name", "📄 Excel – Multiple Names"])


# ---------- Single Name Tab ---------- #
with tab_single:
    st.subheader("Generate Tehillim 119 for a Single Name")

    single_name = st.text_input("Enter a Hebrew name (e.g., יצחק בן אברהם):", value="")

    format_choice_single = st.radio(
        "Choose output format:",
        ["DOCX", "PDF"],
        horizontal=True,
    )

    if st.button("Generate for This Name"):
        if not single_name.strip():
            st.error("Please enter a Hebrew name.")
        else:
            try:
                safe_name = single_name.strip().replace(" ", "_") or "name"

                if format_choice_single == "DOCX":
                    docx_bytes = build_docx_bytes_for_name(single_name, stanzas_119)
                    filename = f"{safe_name}_Tehillim119.docx"
                    st.success("DOCX document generated successfully. Click below to download.")
                    st.download_button(
                        label="⬇ Download DOCX",
                        data=docx_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                else:
                    pdf_bytes = build_pdf_bytes_for_name(single_name, stanzas_119)
                    filename = f"{safe_name}_Tehillim119.pdf"
                    st.success("PDF document generated successfully. Click below to download.")
                    st.download_button(
                        label="⬇ Download PDF",
                        data=pdf_bytes,
                        file_name=filename,
                        mime="application/pdf",
                    )

            except Exception as e:
                st.error(f"Error: {e}")


# ---------- Excel Tab ---------- #
with tab_excel:
    st.subheader("Generate Tehillim 119 for Multiple Names from Excel")

    st.markdown(
        "Upload an Excel file with a column named **`Name`**, "
        "where each row contains one Hebrew name."
    )

    uploaded_file = st.file_uploader(
        "Upload Excel file (.xlsx or .xls):",
        type=["xlsx", "xls"],
    )

    batch_format = st.radio(
        "Output format for each name:",
        ["DOCX", "PDF"],
        horizontal=True,
        key="batch_format_radio",
    )

    if uploaded_file is not None:
        try:
            df = pd.read_excel(uploaded_file)
        except Exception as e:
            st.error(f"Could not read Excel file: {e}")
            df = None

        if df is not None:
            if "Name" not in df.columns:
                st.error("The Excel file must contain a column named 'Name'.")
            else:
                st.write("Preview of names:")
                st.dataframe(df[["Name"]].head())

                if st.button("Generate Files for All Names"):
                    names = [str(n).strip() for n in df["Name"].dropna()]
                    names = [n for n in names if n]

                    if not names:
                        st.error("No valid names found in the 'Name' column.")
                    else:
                        zip_buffer = io.BytesIO()
                        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                            for name in names:
                                try:
                                    safe_name = name.replace(" ", "_") or "name"
                                    if batch_format == "DOCX":
                                        docx_bytes = build_docx_bytes_for_name(name, stanzas_119)
                                        file_name = f"{safe_name}_Tehillim119.docx"
                                        zf.writestr(file_name, docx_bytes)
                                    else:
                                        pdf_bytes = build_pdf_bytes_for_name(name, stanzas_119)
                                        file_name = f"{safe_name}_Tehillim119.pdf"
                                        zf.writestr(file_name, pdf_bytes)
                                except Exception:
                                    # Optionally log per-name errors here
                                    continue
                        zip_buffer.seek(0)

                        st.success(
                            f"Generated {batch_format} files for {len(names)} name(s). "
                            "Download them as a ZIP file below."
                        )
                        st.download_button(
                            label="⬇ Download ZIP of files",
                            data=zip_buffer.getvalue(),
                            file_name=f"Tehillim119_Names_{batch_format}.zip",
                            mime="application/zip",
                        )
    else:
        st.info("Upload an Excel file to enable batch generation.")
